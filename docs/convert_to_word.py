#!/usr/bin/env python3
"""
Convert Android Studio Setup Guide from Markdown to Word Document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph(text, style=style)
    if bold or italic:
        run = p.runs[0]
        run.bold = bold
        run.italic = italic
    return p

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(code)
    p.style = 'Intense Quote'
    run = p.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_bullet_list(doc, items):
    """Add a bullet list"""
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def parse_markdown_to_word(md_file, output_file):
    """Parse markdown and create Word document"""

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title = doc.add_heading('Android Studio Setup Guide for Flutter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('RecallSentry Mobile App - Windows PC')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph()  # Blank line

    # Version info
    version_info = doc.add_paragraph('Document Version: 1.0')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info.runs[0].font.size = Pt(10)

    date_info = doc.add_paragraph('Last Updated: November 2025')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_info.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    in_code_block = False
    code_buffer = []
    in_list = False
    list_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip title and subtitle (already added)
        if i < 5:
            i += 1
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_buffer:
                    add_code_block(doc, '\n'.join(code_buffer))
                    code_buffer = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []
            add_heading(doc, line[5:], level=4)

        # Handle bullet lists
        elif line.strip().startswith(('- ', '* ', '✓ ')):
            list_buffer.append(line.strip()[2:])

        # Handle numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []
            # Extract text after number
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # Handle bold/italic text and regular paragraphs
        elif line.strip():
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []

            # Check for bold
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:  # Odd indices are bold
                        run.bold = True
            else:
                doc.add_paragraph(line.strip())

        # Empty line
        else:
            if list_buffer:
                add_bullet_list(doc, list_buffer)
                list_buffer = []

        i += 1

    # Add any remaining list items
    if list_buffer:
        add_bullet_list(doc, list_buffer)

    # Save document
    doc.save(output_file)
    print(f"SUCCESS: Word document created: {output_file}")

if __name__ == '__main__':
    md_file = 'Android_Studio_Setup_Guide.md'
    output_file = 'Android_Studio_Setup_Guide.docx'

    parse_markdown_to_word(md_file, output_file)
